# -*- coding: utf-8 -*-
"""
Vocab Trainer (DE ↔ FR) – Streamlit Version
- Quizmodi: Multiple Choice & Freitext
- Richtungen: DE→FR, FR→DE
- Import aus Word (.docx): 2-Spalten-Tabelle (DE|FR) oder "de ; fr" pro Zeile
- Mehrere Sammlungen; Tests können auf einzelne Sammlungen gefiltert werden
- Toleranter Vergleich: Groß-/Kleinschreibung + Akzente egal
"""

import json, os, random, unicodedata, io
from dataclasses import dataclass
from typing import List, Dict, Tuple, Optional

import streamlit as st

# Optional: .docx-Import
DOCX_AVAILABLE = True
try:
    from docx import Document  # type: ignore
except Exception:
    DOCX_AVAILABLE = False

APP_DIR = os.path.abspath(os.path.dirname(__file__))
STORE_PATH = os.path.join(APP_DIR, "vocab_store.json")

# -------------------- Helpers --------------------

def normalize(s: str) -> str:
    s = s.strip().lower()
    s = " ".join(s.split())
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    return s

@dataclass
class Entry:
    de: str
    fr: str
    source: str = ""

def load_store() -> Dict:
    if not os.path.exists(STORE_PATH):
        return {"collections": []}
    try:
        with open(STORE_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        # Fallback falls Datei korrupt ist
        return {"collections": []}

def save_store(store: Dict):
    try:
        with open(STORE_PATH, "w", encoding="utf-8") as f:
            json.dump(store, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.warning(f"Konnte Datenbank nicht speichern: {e}")

def get_all_entries(store: Dict) -> List[Entry]:
    out: List[Entry] = []
    for coll in store.get("collections", []):
        src = coll.get("name","?")
        for it in coll.get("items", []):
            out.append(Entry(de=it["de"], fr=it["fr"], source=src))
    return out

def import_docx(file_like, name_hint: str) -> Tuple[str, List[Entry]]:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx ist nicht installiert. Bitte mit 'pip install python-docx' nachrüsten.")
    # file_like: st.UploadedFile oder Bytes
    if hasattr(file_like, "getvalue"):
        data = file_like.getvalue()
        doc = Document(io.BytesIO(data))
    else:
        # Falls Pfad übergeben
        doc = Document(file_like)

    items: List[Entry] = []

    # Tabellen
    for tbl in doc.tables:
        for r_i, row in enumerate(tbl.rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2: 
                continue
            de, fr = cells[0], cells[1]
            if not de or not fr:
                continue
            if r_i == 0 and ("de" in de.lower() and "fr" in fr.lower()):
                # Überschriftenzeile überspringen
                continue
            items.append(Entry(de=de, fr=fr, source=name_hint))

    # Absätze "de ; fr"
    for p in doc.paragraphs:
        t = p.text.strip()
        if ";" in t:
            parts = [s.strip() for s in t.split(";")]
            if len(parts) >= 2 and parts[0] and parts[1]:
                items.append(Entry(de=parts[0], fr=parts[1], source=name_hint))

    # Dedupe
    seen = set(); uniq: List[Entry] = []
    for e in items:
        key = (normalize(e.de), normalize(e.fr))
        if key in seen: 
            continue
        seen.add(key); uniq.append(e)

    coll_name = name_hint or "Import"
    return coll_name, uniq

# -------------------- Built-in Sammlung --------------------

BUILTIN_COLLECTION = {
    "name": "Evolution_und_Steinzeit",
    "items": [
        {"de": "die Urgeschichte", "fr": "la Préhistoire"},
        {"de": "die Frühgeschichte", "fr": "la Protohistoire"},
        {"de": "die Altsteinzeit (2,5 Mio.-9500 v. Chr.)", "fr": "le Paléolithique"},
        {"de": "die Jungsteinzeit (9500 v. Chr.-2200 v. Chr.)", "fr": "le Néolithique"},
        {"de": "der Archäologe", "fr": "l'archéologue"},
        {"de": "die Höhlenmalerei", "fr": "la peinture pariétale"},
        {"de": "der Nomade, die Nomadin", "fr": "un/une nomade"},
        {"de": "roden, urbar machen", "fr": "défricher"},
        {"de": "der/die Sesshafte", "fr": "le/la sédentaire"},
        {"de": "sesshaft werden", "fr": "devenir sédentaire"},
        {"de": "der Tauschhandel", "fr": "le troc"},
        {"de": "der Jäger und Sammler", "fr": "le chasseur-cueilleur"},
        {"de": "der Faustkeil", "fr": "le biface en silex"},
        {"de": "das Haustier", "fr": "l'animal domestique"},
    ],
}

def ensure_builtin(store: Dict) -> Dict:
    names = [c.get("name") for c in store.get("collections", [])]
    if BUILTIN_COLLECTION["name"] not in names:
        store.setdefault("collections", []).append(BUILTIN_COLLECTION)
    return store

# -------------------- Quiz-Helfer --------------------

def qa_pair(e: Entry, mode: str) -> Tuple[str, str]:
    return (e.de, e.fr) if mode == "DE→FR" else (e.fr, e.de)

def build_mc_options(correct: str, session_items: List[Entry], mode: str, all_entries: List[Entry]) -> List[str]:
    # Kandidaten aus Session
    all_answers = list({qa_pair(e, mode)[1] for e in session_items})
    wrongs = [a for a in all_answers if normalize(a) != normalize(correct)]
    random.shuffle(wrongs)
    options = [correct] + wrongs[:3]
    # Falls zu wenig Distraktoren: fülle aus globalen Einträgen
    if len(options) < 4:
        pool_global = [qa_pair(e, mode)[1] for e in all_entries if normalize(qa_pair(e, mode)[1]) != normalize(correct)]
        random.shuffle(pool_global)
        for a in pool_global:
            if a not in options:
                options.append(a)
            if len(options) == 4:
                break
    # Falls immer noch <4, notfalls doppeln
    while len(options) < 4:
        options.append(correct)
    random.shuffle(options)
    return options

# -------------------- Streamlit App --------------------

st.set_page_config(page_title="Vocab Trainer – DE ↔ FR", page_icon="🗂️", layout="wide")

# Initialisierung
if "store" not in st.session_state:
    st.session_state.store = ensure_builtin(load_store())
    save_store(st.session_state.store)

if "quiz" not in st.session_state:
    st.session_state.quiz = None  # wird als Dict verwendet

def reset_quiz():
    st.session_state.quiz = None

# Sidebar
with st.sidebar:
    st.header("⚙️ Einstellungen")
    st.button("Neue Quiz-Session", on_click=reset_quiz)
    st.markdown("---")
    st.subheader("📥 Import (.docx)")
    up = st.file_uploader("Word-Datei hochladen", type=["docx"], help="2-Spalten-Tabelle (DE|FR) oder Zeilen 'de ; fr'")
    default_name = ""
    if up:
        try:
            default_name = os.path.splitext(up.name)[0]
        except Exception:
            default_name = "Import"
    coll_name = st.text_input("Sammlungsname", value=default_name)
    overwrite = st.checkbox("Vorhandene Sammlung gleichen Namens überschreiben", value=False)
    if st.button("Import starten", disabled=(up is None or not coll_name)):
        if up is None:
            st.warning("Bitte eine .docx-Datei auswählen.")
        else:
            try:
                name_hint = coll_name.strip() or default_name or "Import"
                name, items = import_docx(up, name_hint)
                if not items:
                    st.warning("Im Dokument wurden keine Paare erkannt.")
                else:
                    new_coll = {"name": name, "items": [{"de": e.de, "fr": e.fr} for e in items]}
                    st.session_state.store.setdefault("collections", [])
                    names = [c.get("name") for c in st.session_state.store["collections"]]
                    if name in names:
                        if overwrite:
                            idx = names.index(name)
                            st.session_state.store["collections"][idx] = new_coll
                            save_store(st.session_state.store)
                            st.success(f"{len(items)} Einträge in '{name}' importiert (überschrieben).")
                        else:
                            st.error(f"Sammlung '{name}' existiert bereits. Aktiviere 'überschreiben' oder wähle einen anderen Namen.")
                    else:
                        st.session_state.store["collections"].append(new_coll)
                        save_store(st.session_state.store)
                        st.success(f"{len(items)} Einträge in '{name}' importiert.")
                    reset_quiz()
            except Exception as e:
                st.error(f"Import fehlgeschlagen: {e}")

    st.markdown("---")
    st.subheader("📤 Export")
    export_json = json.dumps(st.session_state.store, ensure_ascii=False, indent=2)
    st.download_button("Datenbank als JSON herunterladen", data=export_json.encode("utf-8"),
                       file_name="vocab_store.json", mime="application/json")

    with st.expander("📄 Formatvorlage (.docx)"):
        st.write(
            "- **Tabelle** mit 2 Spalten: Deutsch links, Französisch rechts. Erste Zeile darf Überschriften haben.\n"
            "- **Oder** Textzeilen im Format: `deutsch ; français`\n"
            "- Duplikate werden automatisch entfernt."
        )

# Hauptbereich
st.title("Vocab Trainer – DE ↔ FR")

store = st.session_state.store
all_entries = get_all_entries(store)

# ----------- Startseite (keine aktive Quiz-Session) -----------
if st.session_state.quiz is None:

    cols = st.columns([1.4, 1, 1, 1])
    with cols[0]:
        # Sammlung
        opts = ["(alle)"] + [c.get("name","?") for c in store.get("collections", [])]
        coll = st.selectbox("Sammlung", options=opts, index=0)
    with cols[1]:
        mode = st.radio("Richtung", options=("DE→FR", "FR→DE"), horizontal=False)
    with cols[2]:
        quiztype = st.radio("Quiztyp", options=("Multiple Choice", "Freitext"), horizontal=False)
    with cols[3]:
        num_q = st.number_input("Anzahl Fragen", min_value=5, max_value=100, value=10, step=1)

    # Filter Einträge
    entries = all_entries
    if coll != "(alle)":
        entries = [e for e in entries if e.source == coll]

    if st.button("🎯 Quiz starten", disabled=len(entries) < 4):
        n = min(int(num_q), len(entries))
        chosen = random.sample(entries, n)
        order = list(range(n))
        random.shuffle(order)
        st.session_state.quiz = {
            "items": [{"de": e.de, "fr": e.fr, "source": e.source} for e in chosen],
            "order": order,
            "i": 0,
            "score": 0,
            "history": [],  # (question, given, ok, correct)
            "mode": mode,
            "quiztype": quiztype,
            "phase": "ask",  # "ask" -> "feedback"
            "cached_options": {},  # idx -> options list (MC)
        }

    st.markdown("### 📊 Datenbank")
    if store.get("collections"):
        per = "\n".join([f"- **{c.get('name','?')}**: {len(c.get('items', []))} Einträge" for c in store.get("collections", [])])
        st.write(f"**Gesamt:** {len(all_entries)} Einträge\n\n{per}")
    else:
        st.info("Noch keine externen Sammlungen importiert.")

# ----------- Quiz-Ansicht -----------
else:
    q = st.session_state.quiz
    items = [Entry(**it) for it in q["items"]]
    order = q["order"]
    i = q["i"]
    mode = q["mode"]
    quiztype = q["quiztype"]
    phase = q["phase"]

    if i >= len(order):
        # Fertig
        total = len(order)
        score = q["score"]
        st.success(f"Fertig! Punktzahl: **{score}/{total}** ({round(100*score/total)}%)")
        st.markdown("#### Auswertung")
        # Tabelle
        import pandas as pd
        df = pd.DataFrame(q["history"], columns=["Frage", "Ihre Antwort", "Korrekt", "Richtig"])
        st.dataframe(df, use_container_width=True, hide_index=True)
        st.button("Zur Startseite", on_click=reset_quiz)
    else:
        e = items[order[i]]
        prompt, correct = qa_pair(e, mode)

        topc1, topc2 = st.columns([2,1])
        with topc1:
            st.subheader(f"Quiz – {mode}, {quiztype}")
        with topc2:
            st.write(f"**Frage {i+1}/{len(order)}** • **Punktzahl:** {q['score']}")

        st.markdown(f"**Übersetze:** {prompt}")

        # Eingabe-Bereich
        given_key = f"given_{i}"
        if quiztype == "Multiple Choice":
            # Optionen cachen, damit sie über die Phasen stabil bleiben
            if i not in q["cached_options"]:
                opts = build_mc_options(correct, [items[idx] for idx in order], mode, all_entries)
                q["cached_options"][i] = opts
            else:
                opts = q["cached_options"][i]
            chosen = st.radio("Option wählen", options=opts, index=None, key=given_key)
            st.caption("Tipp: Akzente/Großschreibung egal.")
        else:
            default_val = "" if given_key not in st.session_state else st.session_state[given_key]
            chosen = st.text_input("Antwort eingeben", value=default_val, key=given_key)
            st.caption("Tipp: Akzente/Großschreibung egal. Bestätige mit **Prüfen**.")

        colb1, colb2 = st.columns([1,4])

        # Phase: ask -> prüfen
        if phase == "ask":
            with colb1:
                if st.button("Prüfen"):
                    if not chosen:
                        st.warning("Bitte eine Antwort eingeben/auswählen.")
                    else:
                        ok = normalize(chosen) == normalize(correct)
                        # Feedback speichern
                        q["last_ok"] = ok
                        q["last_given"] = chosen
                        q["phase"] = "feedback"
                        st.rerun()

        # Phase: feedback -> weiter
        elif phase == "feedback":
            ok = q.get("last_ok", False)
            chosen = q.get("last_given", "")
            if ok:
                st.success("✔️ Richtig!")
            else:
                st.error(f"✖️ Falsch. **Richtig:** {correct}")

            # Weiter-Button
            with colb1:
                if st.button("Weiter"):
                    if ok:
                        q["score"] += 1
                    q["history"].append((prompt, chosen, "Ja" if ok else "Nein", correct))
                    q["i"] += 1
                    q["phase"] = "ask"
                    st.rerun()

